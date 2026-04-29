import docx
import sys

def verify(path):
    try:
        doc = docx.Document(path)
        print(f"Number of paragraphs: {len(doc.paragraphs)}")
        print(f"Number of tables: {len(doc.tables)}")
        print(f"First 5 paragraphs:")
        for i in range(min(5, len(doc.paragraphs))):
            print(f"- {doc.paragraphs[i].text[:100]}")
    except Exception as e:
        print(f"Error reading docx: {e}")

if __name__ == "__main__":
    verify(r"c:\Users\user\Desktop\ML-Powered-Scrap-Trading-Platform-main\THE_ULTIMATE_60_PAGE_REPORT.docx")
